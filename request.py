# -*- coding: UTF-8 -*-
import requests
import sys
import json
import time
import math
import getopt
from docx import Document
from docx.shared import Inches, Pt


class WordAndPosition:
    def __init__(self, sentence, pos):
        self.sentence = sentence
        self.pos = pos

    def PixelToSize(self, width, height, angle):
        if angle > 90 or angle < -90:
            print("重拍拜託")
        elif angle > -45 and angle < 45:
            self.height = CalculateDistance(self.pos[0], self.pos[1], self.pos[6], self.pos[7])
            # print(str(self.height) + " " + self.sentence)
        else:
            self.height = CalculateDistance(self.pos[0], self.pos[1], self.pos[2], self.pos[3])
        # 21cm * 30cm
        propotion = float(21 / width)
        self.size = int(propotion * self.height / 0.0358)
        self.indent = int(propotion * self.pos[0] / 0.0358 - 1.5 / 0.0358)
        # print(self.size)


def CalculateDistance(x1, y1, x2, y2):
    return math.sqrt(math.pow(x1 - x2, 2) + math.pow(y1 - y2, 2))

def Usage():
    print('IDK tool')
    print('example:')
    print('python .\\request.py -u {your url}')
    print('or python .\\request.py -f {your file}')
    sys.exit(0)


def CreateWordFile(all_word):
    document = Document()
    for i in all_word:
        paragraph = document.add_paragraph()
        paragraph_format = paragraph.paragraph_format
        paragraph_format.left_indent = Pt(i.indent)
        # paragraph_format.space_before = Pt(18)    # 行間距離
        run = paragraph.add_run(i.sentence)
        font = run.font
        font.size = Pt(i.size)

    document.save('output.docx')

global open_file
global my_files
global is_url
is_url = False
my_files = {}


if not len(sys.argv[1:]):
    Usage()
try:
    opts, args = getopt.getopt(sys.argv[1:],"hu:f:",["help", "url", "file"])
except getopt.GetoptError as err:
    print(str(err))
    Usage()

for o, a in opts:
    if o in ("-h", "--help"):
        Usage()
    elif o in ("-u", "--url"):
        open_file = a
        is_url = True
        my_files = {"url": open_file}
        my_files = json.dumps(my_files)
    elif o in ("-f", "--file"):
        open_file = a
        my_files = open(open_file, 'rb')
        is_url = False

my_endpoint = 'https://westcentralus.api.cognitive.microsoft.com/vision/v3.2-preview.1/read/analyze'
my_params = {}
if is_url:
    content_type = "application/json"
else:
    content_type = "application/octet-stream"
my_headers = {'Content-Type': content_type,
              "Ocp-Apim-Subscription-Key": "e940faf1b4804a5082161f96afd56fd4"}
my_headers2 = {"Ocp-Apim-Subscription-Key": "e940faf1b4804a5082161f96afd56fd4"}
all_word = []

# print(my_files)
r = requests.post(url=my_endpoint, params=my_params,
                  headers=my_headers, data=my_files)
# print(r.url)
content = requests.get(r.headers['Operation-Location'], headers=my_headers2)
content = json.loads(content.content)
while content['status'] != "succeeded":
    time.sleep(5)
    content = requests.get(
        r.headers['Operation-Location'], headers=my_headers2)
    content = json.loads(content.content)
    # print(content['analyzeResult']['readResults'][0]['lines'])

angle = content['analyzeResult']['readResults'][0]['angle']
width = content['analyzeResult']['readResults'][0]['width']
height = content['analyzeResult']['readResults'][0]['height']

for i in content['analyzeResult']['readResults']:
    for j in i['lines']:
        tmp = j['boundingBox']
        all_word.append(WordAndPosition(j['text'], tmp))

for i in all_word:
    i.PixelToSize(width, height, angle)

CreateWordFile(all_word)
print("Creat Word File Over!")
